import json
import hashlib
import subprocess
import sys
from types import SimpleNamespace
from pathlib import Path

import pytest
from pydantic import ValidationError

from doc_web.preview import build_preview
import doc_web.preview_content_hint as preview_content_hint
import doc_web.preview as preview_module
from doc_web.preview_identity import bundle_fingerprint
import doc_web.preview_pdf as preview_pdf
from doc_web.preview_content_hint import build_content_hint
from schemas import (
    DocWebBundleFile,
    DocWebBundleManifest,
    DocWebPreviewMetadata,
    DocWebPreviewSelectorMap,
    DocWebProvenanceBlock,
)


REPO_ROOT = Path(__file__).resolve().parents[1]
BAD_SUMMARY_PHRASES = ("Likely", "likely", "titled or headed")


def _load_jsonl(path: Path) -> list[dict]:
    return [
        json.loads(line)
        for line in path.read_text(encoding="utf-8").splitlines()
        if line.strip()
    ]


def _load_preview(
    bundle_dir: Path,
) -> tuple[DocWebBundleManifest, DocWebPreviewMetadata, list[DocWebProvenanceBlock]]:
    manifest = DocWebBundleManifest(
        **json.loads((bundle_dir / "manifest.json").read_text(encoding="utf-8"))
    )
    metadata = DocWebPreviewMetadata(
        **json.loads((bundle_dir / "preview_metadata.json").read_text(encoding="utf-8"))
    )
    blocks = [
        DocWebProvenanceBlock(**row)
        for row in _load_jsonl(bundle_dir / "provenance" / "blocks.jsonl")
    ]
    return manifest, metadata, blocks


def _assert_direct_summary(summary: str) -> None:
    assert summary
    for phrase in BAD_SUMMARY_PHRASES:
        assert phrase not in summary


def _strings(value):
    if isinstance(value, str):
        yield value
    elif isinstance(value, dict):
        for item in value.values():
            yield from _strings(item)
    elif isinstance(value, list):
        for item in value:
            yield from _strings(item)


def _portable_safe_file_text(bundle_dir: Path) -> str:
    manifest = json.loads((bundle_dir / "manifest.json").read_text(encoding="utf-8"))
    safe_paths = [
        row["path"]
        for row in manifest["files"]
        if row["safe_to_persist"] and row["privacy_class"] == "portable"
    ]
    return "\n".join(
        (bundle_dir / path).read_text(encoding="utf-8") for path in safe_paths
    )


def _cache_identity_fingerprint(payload: dict) -> str:
    source_identity = {
        key: value
        for key, value in payload["source_identity"].items()
        if key != "source_display_label"
    }
    fingerprint_payload = {
        "source_identity": source_identity,
        "doc_web_version": payload["doc_web_version"],
        "doc_web_ref": payload["doc_web_ref"],
        "parser_settings": payload["parser_settings"],
        "runtime_options": payload["runtime_options"],
        "preview_contract_fingerprint": payload["preview_contract_fingerprint"],
        "bundle_fingerprint": payload["bundle_fingerprint"],
        "content_hint": {
            "mode": payload["content_hint"]["mode"],
            "effective_mode": payload["content_hint"]["effective_mode"],
            "provider": payload["content_hint"].get("provider"),
            "model": payload["content_hint"].get("model"),
            "prompt_version": payload["content_hint"].get("prompt_version"),
            "sample_sha256": payload["content_hint"].get("sample_sha256"),
            "cache_key": payload["content_hint"].get("cache_key"),
            "fallback_reason": payload["content_hint"].get("fallback_reason"),
            "requested_timeout_seconds": payload["content_hint"][
                "requested_timeout_seconds"
            ],
        },
    }
    encoded = json.dumps(
        fingerprint_payload,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return f"sha256:{hashlib.sha256(encoded).hexdigest()}"


@pytest.mark.parametrize("module_name", ["doc_web", "doc_web.cli"])
def test_cli_preview_help_is_supported(module_name: str):
    proc = subprocess.run(
        [sys.executable, "-m", module_name, "preview", "--help"],
        cwd=str(REPO_ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )

    assert proc.returncode == 0, proc.stderr
    assert "--input" in proc.stdout
    assert "--out-dir" in proc.stdout
    assert "--content-hint-mode" in proc.stdout


def test_cli_preview_timeout_emits_failed_status(tmp_path: Path):
    bundle_dir = tmp_path / "timeout-preview"
    proc = subprocess.run(
        [
            sys.executable,
            "-m",
            "doc_web",
            "preview",
            "--input",
            "testdata/flat-born-digital-mini.pdf",
            "--out-dir",
            str(bundle_dir),
            "--timeout-seconds",
            "0.001",
            "--content-hint-mode",
            "deterministic",
            "--json",
        ],
        cwd=str(REPO_ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )

    assert proc.returncode == 1
    assert "Traceback" not in proc.stderr
    payload = json.loads(proc.stdout)
    assert payload["status"] == "failed"
    assert payload["status_path"] == "preview_status.jsonl"
    assert not Path(payload["status_path"]).is_absolute()
    assert str(bundle_dir) not in "\n".join(_strings(payload))
    status_rows = _load_jsonl(bundle_dir / "preview_status.jsonl")
    assert status_rows[-1]["stage"] == "failed"
    assert status_rows[-1]["detail"]["reason"] == "timeout"


def test_born_digital_pdf_preview_writes_valid_bundle(tmp_path: Path):
    bundle_dir = tmp_path / "flat-pdf-preview"
    summary = build_preview(
        input_path=REPO_ROOT / "testdata" / "flat-born-digital-mini.pdf",
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )

    manifest, metadata, blocks = _load_preview(bundle_dir)
    selector_map = DocWebPreviewSelectorMap(
        **json.loads(
            (bundle_dir / "preview_to_full_selectors.json").read_text(encoding="utf-8")
        )
    )

    assert summary["coverage_state"] == "complete"
    assert manifest.reading_order == ["page-001", "page-002"]
    assert metadata.coverage_state == "complete"
    assert metadata.structural_facts["text_layer_available"] is True
    assert metadata.content_hint is not None
    assert metadata.content_hint.status == "available"
    _assert_direct_summary(metadata.content_hint.high_level_summary)
    assert metadata.timing_ms["first_status_ms"] <= 500
    assert metadata.timing_ms["preview_ready_ms"] <= 3000
    assert [event.stage for event in metadata.status_events] == [
        "accepted",
        "preparing_pages",
        "detecting_text_or_ocr_need",
        "reading_sample",
        "building_preview_html",
        "preview_ready",
    ]
    assert blocks
    assert all(block.source_page_number in {1, 2} for block in blocks)
    assert all(block.confidence == 1.0 for block in blocks)
    assert selector_map.mappings
    assert all(mapping.mapping_kind == "preserved" for mapping in selector_map.mappings)
    assert (bundle_dir / "cache" / "cache_identity.json").exists()
    parsed_units = _load_jsonl(bundle_dir / "cache" / "parsed_units.jsonl")
    assert len(parsed_units) == len(blocks)
    assert parsed_units[0]["block_id"] == blocks[0].block_id
    assert (
        metadata.cache_identity["reusable_artifacts"]["parsed_units"]
        == "cache/parsed_units.jsonl"
    )
    assert 'id="blk-page-001-0001"' in (bundle_dir / "page-001.html").read_text(
        encoding="utf-8"
    )


def test_mixed_text_image_pdf_fixture_is_reproducible_and_mixed(tmp_path: Path):
    from pypdf import PdfReader

    fixture = REPO_ROOT / "testdata" / "mixed-text-image-mini.pdf"
    source = REPO_ROOT / "testdata" / "mixed-text-image-mini.source.json"
    regenerated = tmp_path / "mixed-text-image-mini.pdf"
    proc = subprocess.run(
        [
            sys.executable,
            "testdata/generate_mixed_text_image_pdf_fixture.py",
            "--output",
            str(regenerated),
        ],
        cwd=str(REPO_ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )

    assert proc.returncode == 0, proc.stderr
    assert (
        hashlib.sha256(regenerated.read_bytes()).hexdigest()
        == hashlib.sha256(fixture.read_bytes()).hexdigest()
    )

    payload = json.loads(source.read_text(encoding="utf-8"))
    assert [page["kind"] for page in payload["pages"]] == ["text-layer", "image-only"]
    reader = PdfReader(str(fixture))
    page_text = [page.extract_text() or "" for page in reader.pages]
    assert len(page_text) == 2
    assert "Page one has usable embedded text" in page_text[0]
    assert page_text[1] == ""


def test_preview_manifest_declares_portable_safe_files(tmp_path: Path):
    bundle_dir = tmp_path / "portable-preview"
    build_preview(
        input_path=REPO_ROOT / "testdata" / "mixed-text-image-mini.pdf",
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )

    manifest = json.loads((bundle_dir / "manifest.json").read_text(encoding="utf-8"))
    safe_files = [row for row in manifest["files"] if row["safe_to_persist"]]
    safe_paths = {row["path"] for row in safe_files}

    assert {
        "manifest.json",
        "index.html",
        "page-001.html",
        "page-002.html",
        "provenance/blocks.jsonl",
        "preview_metadata.json",
        "preview_status.jsonl",
        "preview_to_full_selectors.json",
        "cache/cache_identity.json",
        "cache/parsed_units.jsonl",
    } <= safe_paths
    assert all(not Path(path).is_absolute() for path in safe_paths)
    assert all(".." not in Path(path).parts for path in safe_paths)
    assert all(row["privacy_class"] == "portable" for row in safe_files)
    assert all((bundle_dir / row["path"]).exists() for row in safe_files)
    assert all(
        row["safe_to_replay"] for row in safe_files if row.get("required_for_replay")
    )


@pytest.mark.parametrize(
    "missing_path",
    [
        "manifest.json",
        "index.html",
        "page-001.html",
        "page-002.html",
        "provenance/blocks.jsonl",
        "preview_metadata.json",
        "preview_to_full_selectors.json",
        "cache/cache_identity.json",
        "cache/parsed_units.jsonl",
    ],
)
def test_preview_manifest_schema_rejects_missing_replay_required_files(
    tmp_path: Path, missing_path: str
):
    bundle_dir = tmp_path / "missing-replay-file-preview"
    build_preview(
        input_path=REPO_ROOT / "testdata" / "mixed-text-image-mini.pdf",
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )

    payload = json.loads((bundle_dir / "manifest.json").read_text(encoding="utf-8"))
    payload["files"] = [row for row in payload["files"] if row["path"] != missing_path]

    with pytest.raises(ValidationError, match="missing required replay paths"):
        DocWebBundleManifest(**payload)


@pytest.mark.parametrize(
    "path, wrong_role, error",
    [
        (
            "manifest.json",
            "entry",
            "bundle file manifest role does not match required path",
        ),
        (
            "cache/cache_identity.json",
            "parsed_units",
            "preview bundle file manifest role does not match required path",
        ),
    ],
)
def test_preview_manifest_schema_rejects_required_path_role_mismatch(
    tmp_path: Path, path: str, wrong_role: str, error: str
):
    bundle_dir = tmp_path / "role-mismatch-preview"
    build_preview(
        input_path=REPO_ROOT / "testdata" / "mixed-text-image-mini.pdf",
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )

    payload = json.loads((bundle_dir / "manifest.json").read_text(encoding="utf-8"))
    for row in payload["files"]:
        if row["path"] == path:
            row["role"] = wrong_role
            break
    else:
        raise AssertionError(f"manifest did not include {path}")

    with pytest.raises(ValidationError, match=error):
        DocWebBundleManifest(**payload)


@pytest.mark.parametrize(
    "path",
    [
        "/absolute/manifest.json",
        "../manifest.json",
        r"cache\parsed_units.jsonl",
        "file://manifest.json",
        "s3://bucket/manifest.json",
        "cache//parsed_units.jsonl",
        "C:/manifest.json",
    ],
)
def test_preview_manifest_file_schema_rejects_non_portable_paths(path: str):
    with pytest.raises(ValidationError, match="bundle file paths"):
        DocWebBundleFile(path=path, role="manifest")


@pytest.mark.parametrize(
    "payload, error",
    [
        (
            {
                "path": "manifest.json",
                "role": "manifest",
                "safe_to_persist": True,
                "privacy_class": "private",
            },
            "safe bundle files must use privacy_class='portable'",
        ),
        (
            {
                "path": "manifest.json",
                "role": "manifest",
                "required_for_replay": True,
                "safe_to_replay": False,
                "safe_to_persist": True,
                "privacy_class": "portable",
            },
            "required_for_replay files must be safe_to_replay",
        ),
        (
            {
                "path": "manifest.json",
                "role": "manifest",
                "required_for_replay": True,
                "safe_to_replay": True,
                "safe_to_persist": False,
                "privacy_class": "portable",
            },
            "required_for_replay files must be safe_to_persist",
        ),
    ],
)
def test_preview_manifest_file_schema_rejects_unsafe_replay_claims(
    payload: dict, error: str
):
    with pytest.raises(ValidationError, match=error):
        DocWebBundleFile(**payload)


@pytest.mark.parametrize(
    "role, privacy_class",
    [("debug", "debug"), ("private", "private"), ("cache_local", "cache_local")],
)
def test_preview_manifest_file_schema_rejects_unsafe_roles_marked_safe(
    role: str, privacy_class: str
):
    with pytest.raises(ValidationError, match=f"{role} bundle files must be marked"):
        DocWebBundleFile(
            path=f"cache/{role}.json",
            role=role,
            privacy_class=privacy_class,
            safe_to_persist=True,
        )


@pytest.mark.parametrize(
    "role, privacy_class",
    [("debug", "debug"), ("private", "private"), ("cache_local", "cache_local")],
)
def test_preview_manifest_file_schema_rejects_unsafe_role_portable_privacy(
    role: str, privacy_class: str
):
    with pytest.raises(
        ValidationError,
        match=f"{role} bundle files must use privacy_class='{privacy_class}'",
    ):
        DocWebBundleFile(
            path=f"cache/{role}.json",
            role=role,
            privacy_class="portable",
        )


@pytest.mark.parametrize(
    "source_artifact",
    [
        "/private/client/source.pdf",
        "file:///private/client/source.pdf",
        "s3://private-client/source.pdf",
    ],
)
def test_preview_manifest_schema_rejects_unsafe_source_refs(source_artifact: str):
    replay_file = {
        "safe_to_persist": True,
        "safe_to_replay": True,
        "privacy_class": "portable",
        "required_for_replay": True,
    }
    payload = {
        "document_id": "doc-test",
        "title": "Preview",
        "source_artifact": source_artifact,
        "entries": [
            {
                "entry_id": "page-001",
                "kind": "page",
                "title": "Page 1",
                "path": "page-001.html",
                "order": 1,
                "source_pages": [1],
            }
        ],
        "reading_order": ["page-001"],
        "files": [
            {"path": "manifest.json", "role": "manifest", **replay_file},
            {"path": "index.html", "role": "index", **replay_file},
            {
                "path": "provenance/blocks.jsonl",
                "role": "provenance",
                **replay_file,
            },
            {"path": "page-001.html", "role": "entry", **replay_file},
        ],
    }

    with pytest.raises(ValidationError, match="privacy-safe sha256"):
        DocWebBundleManifest(**payload)


def test_preview_manifest_schema_rejects_preview_module_without_file_manifest():
    payload = {
        "module_id": "doc_web_preview_v1",
        "document_id": "doc-test",
        "title": "Preview",
        "source_artifact": "sha256:" + ("a" * 64),
        "entries": [
            {
                "entry_id": "page-001",
                "kind": "page",
                "title": "Page 1",
                "path": "page-001.html",
                "order": 1,
                "source_pages": [1],
            }
        ],
        "reading_order": ["page-001"],
        "files": [],
    }

    with pytest.raises(
        ValidationError,
        match="preview bundle manifests must include files with required replay paths",
    ):
        DocWebBundleManifest(**payload)


def test_preview_safe_files_do_not_persist_local_source_paths(tmp_path: Path):
    bundle_dir = tmp_path / "private-path-preview"
    source = REPO_ROOT / "testdata" / "mixed-text-image-mini.pdf"
    build_preview(
        input_path=source,
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )

    manifest = json.loads((bundle_dir / "manifest.json").read_text(encoding="utf-8"))
    safe_paths = [
        row["path"]
        for row in manifest["files"]
        if row["safe_to_persist"] and row["privacy_class"] == "portable"
    ]
    persisted_text = "\n".join(
        (bundle_dir / path).read_text(encoding="utf-8") for path in safe_paths
    )

    assert str(REPO_ROOT) not in persisted_text
    assert str(source) not in persisted_text
    assert source.name not in persisted_text
    assert "/var/folders/" not in persisted_text
    assert "/tmp/" not in persisted_text


def test_image_directory_safe_files_do_not_persist_source_filenames(tmp_path: Path):
    bundle_dir = tmp_path / "image-directory-private-path-preview"
    source = REPO_ROOT / "testdata" / "handwritten-notes-mini-images"
    build_preview(
        input_path=source,
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )

    manifest = json.loads((bundle_dir / "manifest.json").read_text(encoding="utf-8"))
    safe_paths = [
        row["path"]
        for row in manifest["files"]
        if row["safe_to_persist"] and row["privacy_class"] == "portable"
    ]
    persisted_text = "\n".join(
        (bundle_dir / path).read_text(encoding="utf-8") for path in safe_paths
    )

    assert str(REPO_ROOT) not in persisted_text
    assert str(source) not in persisted_text
    assert source.name not in persisted_text
    source_image_names = {
        path.name
        for path in source.iterdir()
        if path.suffix.lower() in {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".webp"}
    }
    assert source_image_names
    assert not any(name in persisted_text for name in source_image_names)


def test_image_directory_ocr_error_safe_files_do_not_persist_source_paths(
    tmp_path: Path,
):
    source = tmp_path / "private-client-images"
    source.mkdir()
    private_image = source / "secret-family-page-001.png"
    private_image.write_bytes(b"not actually a png")
    bundle_dir = tmp_path / "image-directory-ocr-error-preview"

    build_preview(
        input_path=source,
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )

    metadata = json.loads(
        (bundle_dir / "preview_metadata.json").read_text(encoding="utf-8")
    )
    errors = metadata["structural_facts"]["ocr_errors"]
    assert errors == [
        {"page": "1", "error": "UnidentifiedImageError: image OCR failed"}
    ]

    manifest = json.loads((bundle_dir / "manifest.json").read_text(encoding="utf-8"))
    safe_paths = [
        row["path"]
        for row in manifest["files"]
        if row["safe_to_persist"] and row["privacy_class"] == "portable"
    ]
    persisted_text = "\n".join(
        (bundle_dir / path).read_text(encoding="utf-8") for path in safe_paths
    )

    assert str(source) not in persisted_text
    assert source.name not in persisted_text
    assert private_image.name not in persisted_text
    assert str(private_image) not in persisted_text


def test_pdf_ocr_error_safe_files_do_not_persist_source_paths(
    tmp_path: Path,
    monkeypatch,
):
    private_marker = "/private/client/secret-family-source.pdf"

    def _fail_ocr(*args, **kwargs):
        raise RuntimeError(f"failed to read {private_marker}")

    monkeypatch.setattr(preview_pdf, "_ocr_pdf_page", _fail_ocr)
    source = REPO_ROOT / "testdata" / "mixed-text-image-mini.pdf"
    bundle_dir = tmp_path / "pdf-ocr-error-preview"

    build_preview(
        input_path=source,
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )

    metadata = json.loads(
        (bundle_dir / "preview_metadata.json").read_text(encoding="utf-8")
    )
    errors = metadata["structural_facts"]["ocr_errors"]
    assert errors == [{"page": "2", "error": "RuntimeError: PDF OCR failed"}]

    manifest = json.loads((bundle_dir / "manifest.json").read_text(encoding="utf-8"))
    safe_paths = [
        row["path"]
        for row in manifest["files"]
        if row["safe_to_persist"] and row["privacy_class"] == "portable"
    ]
    persisted_text = "\n".join(
        (bundle_dir / path).read_text(encoding="utf-8") for path in safe_paths
    )

    assert private_marker not in persisted_text
    assert str(source) not in persisted_text
    assert source.name not in persisted_text
    assert str(REPO_ROOT) not in persisted_text


def test_pdf_metadata_private_paths_do_not_enter_portable_safe_files(tmp_path: Path):
    from pypdf import PdfReader, PdfWriter

    private_marker = "/private/client/secret-family-source.pdf"
    source = REPO_ROOT / "testdata" / "mixed-text-image-mini.pdf"
    private_source = tmp_path / "secret-family-source.pdf"
    reader = PdfReader(str(source))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.add_metadata(
        {
            "/Title": private_marker,
            "/Creator": f"Generated from {private_marker}",
        }
    )
    with private_source.open("wb") as handle:
        writer.write(handle)

    bundle_dir = tmp_path / "pdf-private-metadata-preview"
    build_preview(
        input_path=private_source,
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )

    metadata = json.loads(
        (bundle_dir / "preview_metadata.json").read_text(encoding="utf-8")
    )
    assert metadata["structural_facts"]["metadata_title"] is None
    assert metadata["structural_facts"]["metadata_creator"] is None
    persisted_text = _portable_safe_file_text(bundle_dir)
    assert private_marker not in persisted_text
    assert private_source.name not in persisted_text
    assert str(private_source) not in persisted_text


def test_pdf_metadata_source_stem_does_not_enter_portable_safe_files(tmp_path: Path):
    from pypdf import PdfReader, PdfWriter

    source = REPO_ROOT / "testdata" / "flat-born-digital-mini.pdf"
    private_source = tmp_path / "secret-family-source.pdf"
    reader = PdfReader(str(source))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.add_metadata(
        {
            "/Title": private_source.stem,
            "/Creator": f"Generated from {private_source.stem}",
        }
    )
    with private_source.open("wb") as handle:
        writer.write(handle)

    bundle_dir = tmp_path / "pdf-private-stem-metadata-preview"
    build_preview(
        input_path=private_source,
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )

    metadata = json.loads(
        (bundle_dir / "preview_metadata.json").read_text(encoding="utf-8")
    )
    assert metadata["structural_facts"]["metadata_title"] is None
    assert metadata["structural_facts"]["metadata_creator"] is None
    persisted_text = _portable_safe_file_text(bundle_dir)
    assert private_source.stem not in persisted_text
    assert private_source.name not in persisted_text
    assert str(private_source) not in persisted_text


def test_pdf_metadata_punctuated_donor_filename_does_not_enter_safe_files(
    tmp_path: Path,
):
    from pypdf import PdfReader, PdfWriter

    source = REPO_ROOT / "testdata" / "flat-born-digital-mini.pdf"
    renamed_source = tmp_path / "renamed-upload.pdf"
    reader = PdfReader(str(source))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.add_metadata(
        {
            "/Title": "Internal Report",
            "/Creator": "Generated from client-secret.pdf.",
        }
    )
    with renamed_source.open("wb") as handle:
        writer.write(handle)

    bundle_dir = tmp_path / "pdf-punctuated-donor-metadata-preview"
    build_preview(
        input_path=renamed_source,
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )

    metadata = json.loads(
        (bundle_dir / "preview_metadata.json").read_text(encoding="utf-8")
    )
    assert metadata["structural_facts"]["metadata_creator"] is None
    persisted_text = _portable_safe_file_text(bundle_dir)
    assert "client-secret.pdf" not in persisted_text


def test_pdf_missing_metadata_title_does_not_override_text_hint(tmp_path: Path):
    from pypdf import PdfReader, PdfWriter

    source = REPO_ROOT / "testdata" / "flat-born-digital-mini.pdf"
    no_title_source = tmp_path / "no-title.pdf"
    reader = PdfReader(str(source))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    with no_title_source.open("wb") as handle:
        writer.write(handle)

    bundle_dir = tmp_path / "pdf-no-title-preview"
    build_preview(
        input_path=no_title_source,
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )

    metadata = json.loads(
        (bundle_dir / "preview_metadata.json").read_text(encoding="utf-8")
    )
    assert metadata["structural_facts"]["metadata_title"] is None
    assert metadata["content_hint"]["title_guess"] != "Document Preview"
    assert (
        metadata["content_hint"]["high_level_summary"]
        != "Preview title: Document Preview."
    )


def test_preview_rejects_non_portable_run_id_before_safe_files(tmp_path: Path):
    source = REPO_ROOT / "testdata" / "mixed-text-image-mini.pdf"
    bundle_dir = tmp_path / "unsafe-run-id-preview"

    with pytest.raises(ValueError, match="run_id must be a portable identifier"):
        build_preview(
            input_path=source,
            out_dir=bundle_dir,
            content_hint_mode="deterministic",
            run_id="/private/client/secret-source.pdf",
        )

    assert not (bundle_dir / "manifest.json").exists()
    assert not (bundle_dir / "provenance" / "blocks.jsonl").exists()


def test_preview_rejects_source_filename_run_id_before_safe_files(tmp_path: Path):
    source = REPO_ROOT / "testdata" / "mixed-text-image-mini.pdf"
    bundle_dir = tmp_path / "source-filename-run-id-preview"

    with pytest.raises(ValueError, match="run_id must be a portable identifier"):
        build_preview(
            input_path=source,
            out_dir=bundle_dir,
            content_hint_mode="deterministic",
            run_id="secret-family-source.pdf",
        )

    assert not (bundle_dir / "manifest.json").exists()
    assert not (bundle_dir / "provenance" / "blocks.jsonl").exists()


def test_docx_metadata_private_paths_do_not_enter_portable_safe_files(tmp_path: Path):
    from docx import Document

    private_marker = "/private/client/secret-family-source.docx"
    document = Document(str(REPO_ROOT / "testdata" / "docx-sections-mini.docx"))
    document.core_properties.title = private_marker
    document.core_properties.author = f"Generated from {private_marker}"
    private_source = tmp_path / "secret-family-source.docx"
    document.save(str(private_source))

    bundle_dir = tmp_path / "docx-private-metadata-preview"
    build_preview(
        input_path=private_source,
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )

    metadata = json.loads(
        (bundle_dir / "preview_metadata.json").read_text(encoding="utf-8")
    )
    assert metadata["structural_facts"]["metadata_creator"] is None
    persisted_text = _portable_safe_file_text(bundle_dir)
    assert private_marker not in persisted_text
    assert private_source.name not in persisted_text
    assert str(private_source) not in persisted_text


def test_docx_metadata_source_stem_does_not_leak_as_default_entry_title(
    tmp_path: Path,
):
    from docx import Document

    document = Document()
    private_source = tmp_path / "secret-family-source.docx"
    document.core_properties.title = private_source.stem
    document.core_properties.author = f"Generated from {private_source.stem}"
    document.add_paragraph("Plain paragraph without headings.")
    document.save(str(private_source))

    bundle_dir = tmp_path / "docx-private-stem-default-title-preview"
    build_preview(
        input_path=private_source,
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )

    metadata = json.loads(
        (bundle_dir / "preview_metadata.json").read_text(encoding="utf-8")
    )
    manifest = json.loads((bundle_dir / "manifest.json").read_text(encoding="utf-8"))
    assert metadata["structural_facts"]["metadata_title"] is None
    assert metadata["structural_facts"]["metadata_creator"] is None
    assert manifest["entries"][0]["title"] == "Document Preview"
    persisted_text = _portable_safe_file_text(bundle_dir)
    assert private_source.stem not in persisted_text
    assert private_source.name not in persisted_text
    assert str(private_source) not in persisted_text


def test_docx_title_paragraph_source_stem_does_not_leak_as_entry_title(
    tmp_path: Path,
):
    from docx import Document

    document = Document()
    private_source = tmp_path / "secret-family-source.docx"
    title = document.add_paragraph(private_source.stem)
    title.style = "Title"
    document.add_paragraph("Plain paragraph without headings.")
    document.save(str(private_source))

    bundle_dir = tmp_path / "docx-private-stem-title-paragraph-preview"
    build_preview(
        input_path=private_source,
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )

    metadata = json.loads(
        (bundle_dir / "preview_metadata.json").read_text(encoding="utf-8")
    )
    manifest = json.loads((bundle_dir / "manifest.json").read_text(encoding="utf-8"))
    assert metadata["structural_facts"]["metadata_title"] is None
    assert manifest["entries"][0]["title"] == "Document Preview"
    persisted_text = _portable_safe_file_text(bundle_dir)
    assert private_source.stem not in persisted_text
    assert private_source.name not in persisted_text
    assert str(private_source) not in persisted_text


def test_docx_metadata_title_does_not_leak_as_default_entry_title(tmp_path: Path):
    from docx import Document

    private_marker = "/private/client/secret-family-source.docx"
    document = Document()
    document.core_properties.title = private_marker
    document.add_paragraph("Plain paragraph without headings.")
    private_source = tmp_path / "secret-family-source.docx"
    document.save(str(private_source))

    bundle_dir = tmp_path / "docx-private-default-title-preview"
    build_preview(
        input_path=private_source,
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )

    metadata = json.loads(
        (bundle_dir / "preview_metadata.json").read_text(encoding="utf-8")
    )
    manifest = json.loads((bundle_dir / "manifest.json").read_text(encoding="utf-8"))
    assert metadata["structural_facts"]["metadata_title"] is None
    assert manifest["entries"][0]["title"] == "Document Preview"
    persisted_text = _portable_safe_file_text(bundle_dir)
    assert private_marker not in persisted_text
    assert private_source.name not in persisted_text
    assert str(private_source) not in persisted_text


def test_preview_cache_identity_is_stable_complete_and_privacy_safe(tmp_path: Path):
    source = REPO_ROOT / "testdata" / "mixed-text-image-mini.pdf"
    first_dir = tmp_path / "cache-first"
    second_dir = tmp_path / "cache-second"
    changed_dir = tmp_path / "cache-changed"
    build_preview(
        input_path=source,
        out_dir=first_dir,
        content_hint_mode="deterministic",
    )
    build_preview(
        input_path=source,
        out_dir=second_dir,
        content_hint_mode="deterministic",
    )
    build_preview(
        input_path=source,
        out_dir=changed_dir,
        max_sample_units=1,
        content_hint_mode="deterministic",
    )

    first = json.loads(
        (first_dir / "cache" / "cache_identity.json").read_text(encoding="utf-8")
    )
    second = json.loads(
        (second_dir / "cache" / "cache_identity.json").read_text(encoding="utf-8")
    )
    changed = json.loads(
        (changed_dir / "cache" / "cache_identity.json").read_text(encoding="utf-8")
    )

    assert first["identity_fingerprint"] == second["identity_fingerprint"]
    assert first["identity_fingerprint"] != changed["identity_fingerprint"]
    assert first["source_identity"]["page_count"] == 2
    assert first["source_identity"]["source_sha256"]
    assert first["source_identity"]["source_ref"].startswith("sha256:")
    assert first["doc_web_version"]
    assert first["doc_web_ref"]
    assert first["parser_settings"]["parser"] == "pypdf"
    assert first["parser_settings"]["ocr_rasterizer"] == "pdftoppm-singlefile"
    assert first["bundle_fingerprint"].startswith("sha256:")
    assert "source_artifact" not in first
    assert all(
        not Path(path).is_absolute() for path in first["reusable_artifacts"].values()
    )
    assert source.name not in "\n".join(_strings(first))
    assert str(REPO_ROOT) not in "\n".join(_strings(first))


def test_preview_cache_identity_accepts_integer_content_hint_timeout(tmp_path: Path):
    bundle_dir = tmp_path / "integer-timeout-cache-identity-preview"
    float_bundle_dir = tmp_path / "float-timeout-cache-identity-preview"
    build_preview(
        input_path=REPO_ROOT / "testdata" / "flat-born-digital-mini.pdf",
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
        content_hint_timeout_seconds=1,
    )
    build_preview(
        input_path=REPO_ROOT / "testdata" / "flat-born-digital-mini.pdf",
        out_dir=float_bundle_dir,
        content_hint_mode="deterministic",
        content_hint_timeout_seconds=1.0,
    )

    _load_preview(bundle_dir)
    cache_identity = json.loads(
        (bundle_dir / "cache" / "cache_identity.json").read_text(encoding="utf-8")
    )
    float_cache_identity = json.loads(
        (float_bundle_dir / "cache" / "cache_identity.json").read_text(
            encoding="utf-8"
        )
    )
    assert cache_identity["content_hint"]["requested_timeout_seconds"] == 1.0
    assert cache_identity["runtime_options"]["content_hint_timeout_seconds"] == 1.0
    assert cache_identity["identity_fingerprint"] == _cache_identity_fingerprint(
        cache_identity
    )
    assert cache_identity["identity_fingerprint"] == float_cache_identity[
        "identity_fingerprint"
    ]


def test_cache_identity_schema_cli_rejects_colon_storage_refs_but_allows_sha256(
    tmp_path: Path,
):
    source = REPO_ROOT / "testdata" / "mixed-text-image-mini.pdf"
    bundle_dir = tmp_path / "cache-schema-cli"
    build_preview(
        input_path=source,
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )
    payload = json.loads(
        (bundle_dir / "cache" / "cache_identity.json").read_text(encoding="utf-8")
    )

    allowed_sha_payload = json.loads(json.dumps(payload))
    allowed_sha_payload["runtime_options"]["source_identity_hint"] = (
        "sha256:" + ("c" * 64)
    )
    allowed_sha_payload["identity_fingerprint"] = _cache_identity_fingerprint(
        allowed_sha_payload
    )
    allowed_path = tmp_path / "cache_identity_sha_allowed.json"
    allowed_path.write_text(json.dumps(allowed_sha_payload), encoding="utf-8")
    allowed_proc = subprocess.run(
        [
            sys.executable,
            "validate_artifact.py",
            "--schema",
            "doc_web_cache_identity_v1",
            "--file",
            str(allowed_path),
        ],
        cwd=str(REPO_ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
    )
    assert allowed_proc.returncode == 0, allowed_proc.stdout

    rejected_payload = json.loads(json.dumps(payload))
    rejected_payload["runtime_options"]["source_storage_ref"] = "source:client-pdf"
    rejected_path = tmp_path / "cache_identity_colon_ref.json"
    rejected_path.write_text(json.dumps(rejected_payload), encoding="utf-8")
    rejected_proc = subprocess.run(
        [
            sys.executable,
            "validate_artifact.py",
            "--schema",
            "doc_web_cache_identity_v1",
            "--file",
            str(rejected_path),
        ],
        cwd=str(REPO_ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
    )
    assert rejected_proc.returncode != 0
    assert "cache identity must not contain URI/storage paths" in rejected_proc.stdout

    rejected_source_payload = json.loads(json.dumps(payload))
    rejected_source_payload["source_identity"]["source_ref"] = "source:client-pdf"
    rejected_source_path = tmp_path / "cache_identity_colon_source_ref.json"
    rejected_source_path.write_text(
        json.dumps(rejected_source_payload), encoding="utf-8"
    )
    rejected_source_proc = subprocess.run(
        [
            sys.executable,
            "validate_artifact.py",
            "--schema",
            "doc_web_cache_identity_v1",
            "--file",
            str(rejected_source_path),
        ],
        cwd=str(REPO_ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
    )
    assert rejected_source_proc.returncode != 0
    assert "source_identity.source_ref must be a privacy-safe sha256" in (
        rejected_source_proc.stdout
    )


@pytest.mark.parametrize(
    "field_path, value, error",
    [
        (
            ("preview_contract_fingerprint",),
            "sha256:" + ("a" * 63),
            "cache identity fingerprints must be a sha256:<hex> digest",
        ),
        (
            ("bundle_fingerprint",),
            "sha256:" + ("A" * 64),
            "cache identity fingerprints must be a sha256:<hex> digest",
        ),
        (
            ("identity_fingerprint",),
            "sha256:" + ("d" * 63),
            "cache identity fingerprints must be a sha256:<hex> digest",
        ),
        (
            ("content_hint", "cache_key"),
            "sha256:" + ("b" * 63),
            "content_hint.cache_key must be a sha256:<hex> digest",
        ),
        (
            ("content_hint", "cache_key"),
            "sha256:" + ("B" * 64),
            "content_hint.cache_key must be a sha256:<hex> digest",
        ),
    ],
)
def test_cache_identity_schema_cli_rejects_malformed_sha_refs(
    tmp_path: Path, field_path: tuple[str, ...], value: str, error: str
):
    source = REPO_ROOT / "testdata" / "mixed-text-image-mini.pdf"
    bundle_dir = tmp_path / "cache-schema-malformed-sha"
    build_preview(
        input_path=source,
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )
    payload = json.loads(
        (bundle_dir / "cache" / "cache_identity.json").read_text(encoding="utf-8")
    )
    target = payload
    for key in field_path[:-1]:
        target = target[key]
    target[field_path[-1]] = value
    rejected_path = tmp_path / "cache_identity_malformed_sha.json"
    rejected_path.write_text(json.dumps(payload), encoding="utf-8")

    proc = subprocess.run(
        [
            sys.executable,
            "validate_artifact.py",
            "--schema",
            "doc_web_cache_identity_v1",
            "--file",
            str(rejected_path),
        ],
        cwd=str(REPO_ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
    )

    assert proc.returncode != 0
    assert error in proc.stdout


def test_preview_bundle_fingerprint_ignores_display_labels(tmp_path: Path):
    bundle_dir = tmp_path / "display-label-preview"
    build_preview(
        input_path=REPO_ROOT / "testdata" / "mixed-text-image-mini.pdf",
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )

    manifest = json.loads((bundle_dir / "manifest.json").read_text(encoding="utf-8"))
    blocks = _load_jsonl(bundle_dir / "provenance" / "blocks.jsonl")
    selector_map = json.loads(
        (bundle_dir / "preview_to_full_selectors.json").read_text(encoding="utf-8")
    )
    parsed_units = _load_jsonl(bundle_dir / "cache" / "parsed_units.jsonl")
    before = bundle_fingerprint(
        manifest=manifest,
        provenance_rows=blocks,
        selector_mappings=selector_map["mappings"],
        parsed_units=parsed_units,
    )
    manifest["title"] = "Consumer rewritten display title"
    manifest["entries"][0]["title"] = "Consumer page label"
    after = bundle_fingerprint(
        manifest=manifest,
        provenance_rows=blocks,
        selector_mappings=selector_map["mappings"],
        parsed_units=parsed_units,
    )

    assert before == after


def test_scan_heavy_pdf_preview_runs_bounded_ocr_fallback(tmp_path: Path):
    bundle_dir = tmp_path / "scanned-pdf-preview"
    summary = build_preview(
        input_path=REPO_ROOT / "testdata" / "scanned-prose-mini.pdf",
        out_dir=bundle_dir,
        max_sample_units=1,
        content_hint_mode="deterministic",
    )

    manifest, metadata, blocks = _load_preview(bundle_dir)

    assert summary["coverage_state"] in {"sampled", "complete"}
    assert manifest.reading_order
    assert metadata.coverage_state in {"sampled", "complete"}
    assert metadata.structural_facts["text_layer_available"] is False
    assert metadata.structural_facts["ocr_needed"] is True
    assert metadata.structural_facts["ocr_engine"] == "tesseract"
    assert metadata.structural_facts["ocr_pages"]
    assert metadata.content_hint is not None
    assert metadata.content_hint.status in {"available", "low_quality"}
    _assert_direct_summary(metadata.content_hint.high_level_summary)
    assert blocks
    assert all(block.source_page_number in {1, 2, 3} for block in blocks)
    assert "Preview OCR fallback read" in " ".join(metadata.warnings)


def test_mixed_pdf_preview_ocr_fallback_reads_image_only_page_two(tmp_path: Path):
    bundle_dir = tmp_path / "mixed-pdf-preview"
    build_preview(
        input_path=REPO_ROOT / "testdata" / "mixed-text-image-mini.pdf",
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )

    manifest, metadata, blocks = _load_preview(bundle_dir)

    assert manifest.reading_order == ["page-001", "page-002"]
    assert metadata.coverage_state == "complete"
    assert metadata.structural_facts["page_count"] == 2
    assert metadata.structural_facts["text_layer_pages"] == [1]
    assert metadata.structural_facts["ocr_pages"] == [2]
    assert metadata.structural_facts["ocr_rasterizer"] == "pdftoppm-singlefile"
    page_two_entry = next(
        entry for entry in manifest.entries if entry.entry_id == "page-002"
    )
    assert page_two_entry.source_pages == [2]
    page_two_blocks = [block for block in blocks if block.entry_id == "page-002"]
    assert page_two_blocks
    assert all(block.source_page_number == 2 for block in page_two_blocks)
    assert all(
        source_id.startswith("pdf-page-2-ocr-")
        for block in page_two_blocks
        for source_id in block.source_element_ids
    )
    assert any(
        "IMAGE ONLY PAGE TWO OCR TARGET" in (block.text_quote or "")
        for block in page_two_blocks
    )
    parsed_units = _load_jsonl(bundle_dir / "cache" / "parsed_units.jsonl")
    page_two_units = [unit for unit in parsed_units if unit["entry_id"] == "page-002"]
    assert page_two_units
    assert all(unit["source_page_number"] == 2 for unit in page_two_units)
    assert "IMAGE ONLY PAGE TWO OCR TARGET" in (bundle_dir / "page-002.html").read_text(
        encoding="utf-8"
    )


def test_pdf_raster_missing_output_is_hard_failure(tmp_path: Path, monkeypatch):
    def _fake_run(*args, **kwargs):
        return subprocess.CompletedProcess(args[0], 0, stdout="", stderr="")

    monkeypatch.setattr(preview_pdf.subprocess, "run", _fake_run)
    source = REPO_ROOT / "testdata" / "mixed-text-image-mini.pdf"

    with pytest.raises(
        preview_pdf.PdfRasterizationError,
        match="did not produce expected single-file raster for page 2",
    ) as exc_info:
        preview_pdf._rasterize_pdf_page(source, 2, tmp_path)

    assert str(exc_info.value) == "did not produce expected single-file raster for page 2"
    assert exc_info.value.failure_reason == "pdf_rasterization_missing_output"
    assert str(source) not in str(exc_info.value)
    assert source.name not in str(exc_info.value)


def test_pdf_raster_nonzero_exit_is_hard_failure_and_private(
    tmp_path: Path, monkeypatch
):
    private_marker = "/private/client/secret-family-source.pdf"

    def _fake_run(*args, **kwargs):
        return subprocess.CompletedProcess(
            args[0],
            2,
            stdout="",
            stderr=f"pdftoppm failed for {private_marker}",
        )

    monkeypatch.setattr(preview_pdf.subprocess, "run", _fake_run)
    source = REPO_ROOT / "testdata" / "mixed-text-image-mini.pdf"

    with pytest.raises(
        preview_pdf.PdfRasterizationError,
        match="pdftoppm failed while rasterizing PDF page 2 with exit code 2",
    ) as exc_info:
        preview_pdf._rasterize_pdf_page(source, 2, tmp_path)

    error = str(exc_info.value)
    assert error == "pdftoppm failed while rasterizing PDF page 2 with exit code 2"
    assert exc_info.value.failure_reason == "pdf_rasterization_failed"
    assert private_marker not in error
    assert str(source) not in error
    assert source.name not in error


def test_pdf_raster_timeout_is_hard_failure_and_private(tmp_path: Path, monkeypatch):
    private_marker = "/private/client/secret-family-source.pdf"

    def _fake_run(*args, **kwargs):
        raise subprocess.TimeoutExpired(
            cmd=f"pdftoppm {private_marker}",
            timeout=kwargs.get("timeout"),
            output=f"working on {private_marker}",
            stderr=f"still reading {private_marker}",
        )

    monkeypatch.setattr(preview_pdf.subprocess, "run", _fake_run)
    source = REPO_ROOT / "testdata" / "mixed-text-image-mini.pdf"

    with pytest.raises(
        preview_pdf.PdfRasterizationError,
        match="pdftoppm timed out while rasterizing PDF page 2",
    ) as exc_info:
        preview_pdf._rasterize_pdf_page(source, 2, tmp_path)

    error = str(exc_info.value)
    assert error == "pdftoppm timed out while rasterizing PDF page 2"
    assert exc_info.value.failure_reason == "pdf_rasterization_timeout"
    assert private_marker not in error
    assert str(source) not in error
    assert source.name not in error


def test_pdf_raster_start_failure_is_hard_failure_and_private(
    tmp_path: Path, monkeypatch
):
    private_marker = "/private/client/secret-family-source.pdf"

    def _fake_run(*args, **kwargs):
        raise OSError(f"cannot execute pdftoppm for {private_marker}")

    monkeypatch.setattr(preview_pdf.subprocess, "run", _fake_run)
    source = REPO_ROOT / "testdata" / "mixed-text-image-mini.pdf"

    with pytest.raises(
        preview_pdf.PdfRasterizationError,
        match="pdftoppm could not be started while rasterizing PDF page 2",
    ) as exc_info:
        preview_pdf._rasterize_pdf_page(source, 2, tmp_path)

    error = str(exc_info.value)
    assert error == "pdftoppm could not be started while rasterizing PDF page 2"
    assert exc_info.value.failure_reason == "pdf_rasterization_start_failed"
    assert private_marker not in error
    assert str(source) not in error
    assert source.name not in error


@pytest.mark.parametrize(
    "mode, expected_reason, expected_error",
    [
        (
            "missing",
            "pdf_rasterization_missing_output",
            "did not produce expected single-file raster for page 2",
        ),
        (
            "nonzero",
            "pdf_rasterization_failed",
            "pdftoppm failed while rasterizing PDF page 2 with exit code 2",
        ),
        (
            "timeout",
            "pdf_rasterization_timeout",
            "pdftoppm timed out while rasterizing PDF page 2",
        ),
        (
            "start",
            "pdf_rasterization_start_failed",
            "pdftoppm could not be started while rasterizing PDF page 2",
        ),
    ],
)
def test_cli_preview_pdf_raster_failures_emit_failed_status_safely(
    tmp_path: Path,
    monkeypatch,
    capsys,
    mode: str,
    expected_reason: str,
    expected_error: str,
):
    private_marker = "/private/client/secret-family-source.pdf"

    def _fake_run(*args, **kwargs):
        if mode == "missing":
            return subprocess.CompletedProcess(args[0], 0, stdout="", stderr="")
        if mode == "nonzero":
            return subprocess.CompletedProcess(
                args[0],
                2,
                stdout="",
                stderr=f"pdftoppm failed for {private_marker}",
            )
        if mode == "timeout":
            raise subprocess.TimeoutExpired(
                cmd=f"pdftoppm {private_marker}",
                timeout=kwargs.get("timeout"),
                output=f"working on {private_marker}",
                stderr=f"still reading {private_marker}",
            )
        raise OSError(f"cannot execute pdftoppm for {private_marker}")

    monkeypatch.setattr(preview_pdf.subprocess, "run", _fake_run)
    source = REPO_ROOT / "testdata" / "mixed-text-image-mini.pdf"
    bundle_dir = tmp_path / f"pdf-raster-{mode}-preview"
    monkeypatch.setattr(
        sys,
        "argv",
        [
            "doc-web",
            "preview",
            "--input",
            str(source),
            "--out-dir",
            str(bundle_dir),
            "--content-hint-mode",
            "deterministic",
            "--json",
        ],
    )
    from doc_web.cli import main as cli_main

    with pytest.raises(SystemExit) as exc_info:
        cli_main()

    assert exc_info.value.code == 1
    captured = capsys.readouterr()
    assert "Traceback" not in captured.out
    assert "Traceback" not in captured.err
    payload = json.loads(captured.out)
    assert payload["status"] == "failed"
    assert payload["error"] == expected_error
    assert payload["status_path"] == "preview_status.jsonl"
    assert not Path(payload["status_path"]).is_absolute()

    status_rows = _load_jsonl(bundle_dir / "preview_status.jsonl")
    assert status_rows[-1]["stage"] == "failed"
    assert status_rows[-1]["message"] == expected_error
    assert status_rows[-1]["detail"]["reason"] == expected_reason
    combined_output = "\n".join(
        [
            captured.out,
            captured.err,
            json.dumps(status_rows, ensure_ascii=False),
        ]
    )
    assert private_marker not in combined_output
    assert str(source) not in combined_output
    assert source.name not in combined_output
    assert str(bundle_dir) not in combined_output


def test_pdf_preview_deferred_skipped_units_are_unique_after_empty_ocr(
    tmp_path: Path,
    monkeypatch,
):
    monkeypatch.setattr(preview_pdf, "_ocr_pdf_page", lambda *args, **kwargs: "")
    bundle_dir = tmp_path / "pdf-empty-ocr-preview"

    build_preview(
        input_path=REPO_ROOT / "testdata" / "scanned-prose-mini.pdf",
        out_dir=bundle_dir,
        max_sample_units=1,
        content_hint_mode="deterministic",
    )

    metadata = json.loads(
        (bundle_dir / "preview_metadata.json").read_text(encoding="utf-8")
    )
    skipped_pages = [
        unit for unit in metadata["skipped_units"] if unit["kind"] == "page"
    ]
    skipped_identifiers = [unit["identifier"] for unit in skipped_pages]
    page_count = metadata["structural_facts"]["page_count"]

    assert len(skipped_identifiers) == len(set(skipped_identifiers))
    assert sorted(skipped_identifiers, key=int) == [
        str(page_index) for page_index in range(1, page_count + 1)
    ]
    reasons = {unit["identifier"]: unit["reason"] for unit in skipped_pages}
    assert reasons["1"] == "no_preview_ocr_text"
    assert set(reasons.values()) <= {"no_preview_ocr_text", "no_text_layer"}


def test_image_directory_preview_runs_bounded_ocr(tmp_path: Path):
    bundle_dir = tmp_path / "image-directory-preview"
    summary = build_preview(
        input_path=REPO_ROOT / "testdata" / "handwritten-notes-mini-images",
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )

    manifest, metadata, blocks = _load_preview(bundle_dir)

    assert summary["coverage_state"] == "sampled"
    assert manifest.reading_order == ["page-002"]
    assert metadata.coverage_state == "sampled"
    assert metadata.structural_facts["format"] == "image_directory"
    assert metadata.structural_facts["image_count"] == 2
    assert metadata.structural_facts["sampled_image_count"] == 1
    assert metadata.structural_facts["ocr_engine"] == "tesseract"
    assert metadata.structural_facts["ocr_text_chars"] > 50
    assert metadata.structural_facts["text_layer_available"] is False
    assert metadata.structural_facts["ocr_needed"] is True
    assert metadata.skipped_units
    assert blocks
    assert metadata.content_hint is not None
    assert metadata.content_hint.status == "available"
    _assert_direct_summary(metadata.content_hint.high_level_summary)
    parsed_units = _load_jsonl(bundle_dir / "cache" / "parsed_units.jsonl")
    assert len(parsed_units) == len(blocks)
    html = (bundle_dir / "page-002.html").read_text(encoding="utf-8")
    assert "attic boxes" in html


def test_content_hint_summary_is_direct_and_strips_source_site():
    hint = build_content_hint(
        facts={
            "format": "pdf",
            "metadata_title": "Robo Rally Rulebook - 1jour-1jeu.com",
            "text_layer_available": True,
        },
        parsed_units=[
            {
                "text": (
                    "Game contents include robot figures, reboot tokens, "
                    "programming cards, checkpoints, and gameboards."
                )
            }
        ],
        coverage_state="sampled",
        warnings=[],
    )

    assert hint["status"] == "available"
    assert hint["title_guess"] == "Robo Rally Rulebook"
    assert hint["document_kind_hint"] == "game rulebook"
    assert hint["high_level_summary"] == (
        "Robo Rally Rulebook is a game rulebook covering components and "
        "robot programming."
    )
    _assert_direct_summary(hint["high_level_summary"])
    assert " or " not in hint["document_kind_hint"]


def test_content_hint_deterministic_family_history_title_summary():
    hint = build_content_hint(
        facts={"format": "image_directory"},
        parsed_units=[
            {"text": ("ONWARD TO THE UNKNOWN 1887 - 1987 Moise and Sophie L’Heureux")}
        ],
        coverage_state="sampled",
        warnings=[],
        mode="deterministic",
    )

    assert hint["document_kind_hint"] == "family history"
    assert hint["high_level_summary"] == (
        "ONWARD TO THE UNKNOWN 1887-1987 is a family history about "
        "Moise and Sophie L’Heureux."
    )


def test_content_hint_deterministic_family_subtitle_summary():
    hint = build_content_hint(
        facts={
            "format": "pdf",
            "metadata_title": (
                "Onward to the Unknown: A Genealogy and Biography of the "
                "L'Heureux Family"
            ),
        },
        parsed_units=[{"text": "genealogy biography family descendants"}],
        coverage_state="sampled",
        warnings=[],
        mode="deterministic",
    )

    assert hint["document_kind_hint"] == "family history"
    assert hint["high_level_summary"] == (
        "Onward to the Unknown is a family history about the L'Heureux family."
    )


def test_content_hint_ai_pass_uses_model_json(monkeypatch):
    calls: list[dict] = []
    clients: list[dict] = []

    class _DummyCompletions:
        def create(self, **kwargs):
            calls.append(kwargs)
            return SimpleNamespace(
                choices=[
                    SimpleNamespace(
                        message=SimpleNamespace(
                            content=json.dumps(
                                {
                                    "title_guess": "Onward to the Unknown",
                                    "document_kind_hint": "family history",
                                    "high_level_summary": (
                                        "The document traces the L'Heureux "
                                        "family history."
                                    ),
                                }
                            )
                        )
                    )
                ],
                usage=SimpleNamespace(prompt_tokens=100, completion_tokens=30),
            )

    def _dummy_openai_client(**kwargs):
        clients.append(kwargs)
        return SimpleNamespace(chat=SimpleNamespace(completions=_DummyCompletions()))

    monkeypatch.setattr(
        preview_content_hint, "_make_openai_client", _dummy_openai_client
    )
    monkeypatch.setattr(
        preview_content_hint,
        "get_doc_web_api_key",
        lambda provider, env=None: "test-key",
    )

    hint = build_content_hint(
        facts={
            "format": "pdf",
            "metadata_title": (
                "Onward to the Unknown: A Genealogy and Biography of the "
                "L'Heureux Family"
            ),
            "metadata_creator": "OCRmyPDF 16.6.2 / Tesseract OCR-hOCR 5.5.0",
        },
        parsed_units=[
            {
                "text": (
                    "Moise and Sophie L'Heureux family biography genealogy "
                    "children born married descendants."
                )
            }
        ],
        coverage_state="sampled",
        warnings=[],
        mode="ai",
        ai_model="test-cheap-model",
        ai_timeout_seconds=0.5,
        source_sha256="a" * 64,
    )

    assert hint["summary_provider"] == "openai"
    assert hint["summary_model"] == "test-cheap-model"
    assert hint["document_kind_hint"] == "family history"
    assert hint["high_level_summary"] == (
        "Onward to the Unknown traces the L'Heureux family history."
    )
    assert "ai_summary:openai:test-cheap-model" in hint["basis"]
    assert hint["cache_key"].startswith("sha256:")
    assert clients[0]["max_retries"] == 0
    assert calls[0]["model"] == "test-cheap-model"
    assert calls[0]["response_format"] == {"type": "json_object"}


def test_content_hint_ai_trims_long_title_prefix(monkeypatch):
    class _DummyCompletions:
        def create(self, **kwargs):
            title = (
                "Onward to the Unknown: A Genealogy and Biography of the "
                "L'Heureux Family"
            )
            return SimpleNamespace(
                choices=[
                    SimpleNamespace(
                        message=SimpleNamespace(
                            content=json.dumps(
                                {
                                    "title_guess": title,
                                    "document_kind_hint": "family history",
                                    "high_level_summary": (
                                        f"{title} is a family history detailing "
                                        "the genealogy and biography of the "
                                        "L'Heureux family."
                                    ),
                                }
                            )
                        )
                    )
                ],
                usage=SimpleNamespace(prompt_tokens=100, completion_tokens=30),
            )

    def _dummy_openai_client(**kwargs):
        return SimpleNamespace(chat=SimpleNamespace(completions=_DummyCompletions()))

    monkeypatch.setattr(
        preview_content_hint, "_make_openai_client", _dummy_openai_client
    )
    monkeypatch.setattr(
        preview_content_hint,
        "get_doc_web_api_key",
        lambda provider, env=None: "test-key",
    )

    hint = build_content_hint(
        facts={"format": "pdf"},
        parsed_units=[{"text": "family genealogy biography descendants"}],
        coverage_state="sampled",
        warnings=[],
        mode="ai",
        source_sha256="b" * 64,
    )

    assert hint["high_level_summary"] == (
        "A family history detailing the genealogy and biography of the "
        "L'Heureux family."
    )


def test_auto_content_hint_respects_preview_deadline(tmp_path: Path, monkeypatch):
    calls: list[dict] = []

    def _fake_content_hint(**kwargs):
        calls.append(kwargs)
        return {
            "status": "available",
            "title_guess": "Acme Community Arts Initiative",
            "document_kind_hint": "report",
            "high_level_summary": "Acme Community Arts Initiative is a report.",
            "basis": ["test"],
            "evidence": [],
            "warnings": [],
            "text_quality_score": 1.0,
            "coverage_state": "complete",
        }

    monkeypatch.setattr(preview_module, "build_content_hint", _fake_content_hint)

    build_preview(
        input_path=REPO_ROOT / "testdata" / "flat-born-digital-mini.pdf",
        out_dir=tmp_path / "deadline-preview",
        usable_deadline_seconds=0.001,
        content_hint_mode="auto",
    )

    assert calls[0]["mode"] == "deterministic"


def test_auto_content_hint_skips_ai_for_image_directory(tmp_path: Path, monkeypatch):
    calls: list[dict] = []

    def _fake_content_hint(**kwargs):
        calls.append(kwargs)
        return {
            "status": "available",
            "title_guess": "Handwritten Notes",
            "document_kind_hint": "notes",
            "high_level_summary": "Handwritten Notes is a notes collection.",
            "basis": ["test"],
            "evidence": [],
            "warnings": [],
            "text_quality_score": 1.0,
            "coverage_state": "sampled",
        }

    monkeypatch.setattr(preview_module, "build_content_hint", _fake_content_hint)

    build_preview(
        input_path=REPO_ROOT / "testdata" / "handwritten-notes-mini-images",
        out_dir=tmp_path / "image-auto-preview",
        content_hint_mode="auto",
    )

    assert calls[0]["mode"] == "deterministic"


def test_auto_content_hint_cache_identity_records_resolved_effective_mode(
    tmp_path: Path, monkeypatch
):
    monkeypatch.setattr(
        preview_content_hint,
        "get_doc_web_api_key",
        lambda provider, env=None: None,
    )

    bundle_dir = tmp_path / "auto-cache-identity-preview"
    build_preview(
        input_path=REPO_ROOT / "testdata" / "flat-born-digital-mini.pdf",
        out_dir=bundle_dir,
    )

    metadata = json.loads(
        (bundle_dir / "preview_metadata.json").read_text(encoding="utf-8")
    )
    cache_hint = metadata["cache_identity"]["content_hint"]
    assert cache_hint["mode"] == "auto"
    assert cache_hint["effective_mode"] == "deterministic"
    assert cache_hint["fallback_reason"] == "DOC_WEB_OPENAI_API_KEY is not configured."


def test_ai_content_hint_fallback_reason_is_portable_in_safe_metadata(
    tmp_path: Path, monkeypatch
):
    private_marker = "/private/client/secret-family-source.pdf"

    def _raise_openai_client(**kwargs):
        raise TimeoutError(f"while reading {private_marker}")

    monkeypatch.setattr(
        preview_content_hint,
        "get_doc_web_api_key",
        lambda provider, env=None: "test-key",
    )
    monkeypatch.setattr(
        preview_content_hint,
        "_make_openai_client",
        _raise_openai_client,
    )

    bundle_dir = tmp_path / "ai-fallback-cache-identity-preview"
    build_preview(
        input_path=REPO_ROOT / "testdata" / "flat-born-digital-mini.pdf",
        out_dir=bundle_dir,
        content_hint_mode="ai",
    )

    metadata_text = (bundle_dir / "preview_metadata.json").read_text(encoding="utf-8")
    metadata = json.loads(metadata_text)
    cache_hint = metadata["cache_identity"]["content_hint"]
    assert cache_hint["mode"] == "ai"
    assert cache_hint["effective_mode"] == "deterministic"
    assert cache_hint["fallback_reason"] == "TimeoutError content hint fallback"
    assert private_marker not in metadata_text
    assert "TimeoutError:" not in metadata_text


def test_preview_cache_identity_allows_colon_model_ids(tmp_path: Path):
    model_id = "ft:gpt-4.1-nano:org:custom:abc"
    bundle_dir = tmp_path / "fine-tuned-model-cache-identity-preview"

    build_preview(
        input_path=REPO_ROOT / "testdata" / "flat-born-digital-mini.pdf",
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
        content_hint_model=model_id,
    )

    cache_identity = json.loads(
        (bundle_dir / "cache" / "cache_identity.json").read_text(encoding="utf-8")
    )
    metadata = json.loads(
        (bundle_dir / "preview_metadata.json").read_text(encoding="utf-8")
    )

    assert cache_identity["runtime_options"]["content_hint_model"] == model_id
    assert cache_identity["content_hint"]["model"] == model_id
    assert metadata["cache_identity"]["content_hint"]["model"] == model_id


def test_docx_preview_is_pageless_and_selector_stable(tmp_path: Path):
    bundle_dir = tmp_path / "docx-preview"
    summary = build_preview(
        input_path=REPO_ROOT / "testdata" / "docx-sections-mini.docx",
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )

    manifest, metadata, blocks = _load_preview(bundle_dir)
    selector_map = DocWebPreviewSelectorMap(
        **json.loads(
            (bundle_dir / "preview_to_full_selectors.json").read_text(encoding="utf-8")
        )
    )

    assert summary["coverage_state"] == "complete"
    assert manifest.reading_order == ["chapter-001", "chapter-002"]
    assert [entry.title for entry in manifest.entries] == ["Overview", "Roster"]
    assert metadata.structural_facts["pageless"] is True
    assert metadata.structural_facts["paragraph_count"] == 7
    assert blocks
    assert all(block.source_page_number is None for block in blocks)
    assert all(block.source_element_ids[0].startswith("docx-") for block in blocks)
    assert selector_map.mappings
    assert all(
        mapping.preview_block_id == mapping.full_block_id
        for mapping in selector_map.mappings
    )
    parsed_units = _load_jsonl(bundle_dir / "cache" / "parsed_units.jsonl")
    assert parsed_units[2]["source_element_ids"] == ["docx-paragraph-0004"]
    assert "The wider proof should keep both paragraphs" in (
        bundle_dir / "chapter-001.html"
    ).read_text(encoding="utf-8")


def test_preview_artifacts_validate_with_schema_cli(tmp_path: Path):
    bundle_dir = tmp_path / "cli-preview"
    build_preview(
        input_path=REPO_ROOT / "testdata" / "flat-born-digital-mini.pdf",
        out_dir=bundle_dir,
        content_hint_mode="deterministic",
    )

    for schema, filename in [
        ("doc_web_bundle_manifest_v1", "manifest.json"),
        ("doc_web_provenance_block_v1", "provenance/blocks.jsonl"),
        ("doc_web_preview_metadata_v1", "preview_metadata.json"),
        ("doc_web_preview_selector_map_v1", "preview_to_full_selectors.json"),
        ("doc_web_cache_identity_v1", "cache/cache_identity.json"),
    ]:
        proc = subprocess.run(
            [
                sys.executable,
                "validate_artifact.py",
                "--schema",
                schema,
                "--file",
                str(bundle_dir / filename),
            ],
            cwd=str(REPO_ROOT),
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
        )
        assert proc.returncode == 0, proc.stdout
