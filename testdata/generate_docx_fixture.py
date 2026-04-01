#!/usr/bin/env python3
"""Generate the repo-owned DOCX fixture from checked-in source data."""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent
DEFAULT_SOURCE = ROOT / "docx-mini.source.json"
DEFAULT_OUTPUT = ROOT / "docx-mini.docx"


def _add_table(document: Document, table_payload: dict) -> None:
    headers = table_payload["headers"]
    rows = table_payload["rows"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)


def _render_block(document: Document, block: dict) -> None:
    block_type = block["type"]
    if block_type == "heading":
        document.add_heading(block["text"], level=int(block.get("level", 2)))
        return
    if block_type == "paragraph":
        document.add_paragraph(block["text"])
        return
    if block_type == "bullet":
        document.add_paragraph(block["text"], style="List Bullet")
        return
    if block_type == "table":
        _add_table(document, block)
        return
    raise ValueError(f"Unsupported DOCX fixture block type: {block_type}")


def build_fixture(source_path: Path, output_path: Path) -> None:
    payload = json.loads(source_path.read_text(encoding="utf-8"))
    document = Document()
    document.add_heading(payload["title"], level=0)

    for section in payload["sections"]:
        document.add_heading(section["heading"], level=1)
        if section.get("blocks"):
            for block in section["blocks"]:
                _render_block(document, block)
            continue

        for paragraph in section.get("paragraphs", []):
            document.add_paragraph(paragraph)
        for bullet in section.get("bullets", []):
            document.add_paragraph(bullet, style="List Bullet")
        table_payload = section.get("table")
        if table_payload:
            _add_table(document, table_payload)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE, help="Checked-in source JSON")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT, help="Output DOCX path")
    args = parser.parse_args()

    build_fixture(args.source, args.output)
    print(f"Wrote {args.output}")


if __name__ == "__main__":
    main()
